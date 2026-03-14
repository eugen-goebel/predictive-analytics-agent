"""
Report Generator — Creates a professional DOCX report with ML results.

Combines all pipeline results into a formatted Word document with tables,
charts, and auto-generated conclusions.
"""

import os
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from agents.data_profiler import DataProfile
from agents.preprocessor import PreprocessResult
from agents.feature_engineer import FeatureResult
from agents.model_trainer import TrainingResult
from agents.evaluator import EvaluationResult


# ---------------------------------------------------------------------------
# Colors
# ---------------------------------------------------------------------------

COLOR_DARK = RGBColor(0x1B, 0x2A, 0x4A)
COLOR_ACCENT = RGBColor(0x0D, 0x73, 0x77)
COLOR_LIGHT_GRAY = "F5F5F5"
COLOR_WHITE = "FFFFFF"
COLOR_HEADER_BG = "1B2A4A"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, color_hex: str):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(
        qn("w:shd"),
        {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex},
    )
    shading.append(shading_elm)


def _add_heading(doc, text: str, level: int = 1):
    """Add a colored heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = COLOR_DARK


def _add_table(doc, headers: list[str], rows: list[list[str]], highlight_col: int = -1):
    """Add a styled table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        _set_cell_bg(cell, COLOR_HEADER_BG)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row in enumerate(rows):
        bg = COLOR_LIGHT_GRAY if r_idx % 2 == 0 else COLOR_WHITE
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            _set_cell_bg(cell, bg)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    return table


# ---------------------------------------------------------------------------
# Main function
# ---------------------------------------------------------------------------

def generate_docx_report(
    profile: DataProfile,
    preprocess_result: PreprocessResult,
    feature_result: FeatureResult,
    training_result: TrainingResult,
    eval_result: EvaluationResult,
    output_dir: str = "output",
) -> str:
    """
    Generate a professional DOCX report from pipeline results.

    Returns:
        Absolute path to the generated .docx file
    """
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()

    # --- Cover Page ---
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Predictive Analytics Report")
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_DARK
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run(f"Dataset: {profile.filename}")
    run2.font.size = Pt(14)
    run2.font.color.rgb = COLOR_ACCENT

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = date_para.add_run(f"Generated: {date.today().strftime('%B %d, %Y')}")
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = tagline.add_run("Automated Machine Learning Pipeline")
    run4.font.size = Pt(10)
    run4.font.color.rgb = COLOR_ACCENT

    doc.add_page_break()

    # --- 1. Data Profile ---
    _add_heading(doc, "1. Data Profile")
    task_label = profile.task_type.capitalize()
    _add_table(doc,
        ["Property", "Value"],
        [
            ["Dataset", profile.filename],
            ["Rows", f"{profile.row_count:,}"],
            ["Columns", str(profile.column_count)],
            ["Target Column", profile.target_column],
            ["Task Type", task_label],
            ["Data Quality", f"{profile.data_quality_score}/100"],
        ],
    )
    doc.add_paragraph()

    # Class distribution
    if profile.class_distribution:
        _add_heading(doc, "Class Distribution", level=2)
        rows = [[cls, str(count)] for cls, count in profile.class_distribution.items()]
        _add_table(doc, ["Class", "Count"], rows)
        doc.add_paragraph()

    # --- 2. Column Overview ---
    _add_heading(doc, "2. Column Overview")
    col_rows = [
        [cp.name, cp.dtype, str(cp.missing_count), str(cp.unique_count)]
        for cp in profile.column_profiles
    ]
    _add_table(doc, ["Column", "Type", "Missing", "Unique Values"], col_rows)
    doc.add_paragraph()

    # --- 3. Preprocessing ---
    _add_heading(doc, "3. Preprocessing Steps")
    for step in preprocess_result.steps_applied:
        doc.add_paragraph(step, style="List Bullet")
    doc.add_paragraph()

    # --- 4. Feature Selection ---
    _add_heading(doc, "4. Feature Selection")
    doc.add_paragraph(
        f"Method: {feature_result.method}\n"
        f"Selected {feature_result.selected_feature_count} of "
        f"{feature_result.original_feature_count} features."
    )

    if feature_result.feature_importances:
        imp_rows = [
            [fi.name, f"{fi.importance:.2f}"]
            for fi in feature_result.feature_importances[:10]
        ]
        _add_table(doc, ["Feature", "Importance Score"], imp_rows)
    doc.add_paragraph()

    if feature_result.dropped_features:
        doc.add_paragraph(
            f"Dropped features: {', '.join(feature_result.dropped_features)}"
        )
        doc.add_paragraph()

    # --- 5. Model Comparison ---
    _add_heading(doc, "5. Model Comparison")
    metric_name = "Accuracy" if training_result.task_type == "classification" else "R² Score"
    model_rows = [
        [
            ("*** " + ms.name + " ***") if ms.name == training_result.best_model_name else ms.name,
            f"{ms.score:.4f}",
            f"±{ms.cv_std:.4f}",
            f"{ms.training_time:.3f}s",
        ]
        for ms in training_result.model_scores
    ]
    _add_table(doc, ["Model", metric_name, "Std Dev", "Time"], model_rows)
    doc.add_paragraph(
        f"Best model: {training_result.best_model_name} "
        f"(score: {training_result.best_score:.4f})"
    )
    doc.add_paragraph()

    # --- 6. Evaluation ---
    _add_heading(doc, "6. Model Evaluation")
    eval_rows = [
        ["Test Score", f"{eval_result.test_score:.4f}"],
        ["Train Score", f"{eval_result.train_score:.4f}"],
        ["Overfitting", "Yes — consider simplifying the model" if eval_result.is_overfitting else "No"],
    ]
    for key, val in eval_result.metrics.items():
        eval_rows.append([key.replace("_", " ").title(), f"{val:.4f}" if isinstance(val, float) else str(val)])
    _add_table(doc, ["Metric", "Value"], eval_rows)
    doc.add_paragraph()

    # Classification report
    if eval_result.classification_report:
        _add_heading(doc, "Classification Report", level=2)
        p = doc.add_paragraph()
        run = p.add_run(eval_result.classification_report)
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        doc.add_paragraph()

    # --- 7. Charts ---
    _add_heading(doc, "7. Visualizations")
    chart_dir = os.path.join(output_dir, "charts")
    for chart in eval_result.charts:
        chart_path = os.path.join(chart_dir, chart.filename)
        if os.path.exists(chart_path):
            doc.add_paragraph(chart.title, style="Heading 3")
            doc.add_picture(chart_path, width=Inches(5.5))
            doc.add_paragraph(chart.description)
            doc.add_paragraph()

    # --- 8. Conclusion ---
    doc.add_page_break()
    _add_heading(doc, "8. Conclusion")

    score_label = "accuracy" if training_result.task_type == "classification" else "R² score"
    conclusion = (
        f"The automated ML pipeline analyzed the '{profile.filename}' dataset "
        f"({profile.row_count:,} rows, {profile.column_count} columns) and identified "
        f"'{profile.target_column}' as the prediction target ({profile.task_type}).\n\n"
        f"After preprocessing, {feature_result.selected_feature_count} features were selected "
        f"from {feature_result.original_feature_count} using {feature_result.method}.\n\n"
        f"Four models were compared using 5-fold cross-validation. "
        f"The best performer was {training_result.best_model_name} with a "
        f"{score_label} of {eval_result.test_score:.4f} on the holdout test set."
    )

    if eval_result.is_overfitting:
        conclusion += (
            f"\n\nNote: The model shows signs of overfitting "
            f"(train: {eval_result.train_score:.4f}, test: {eval_result.test_score:.4f}). "
            f"Consider using regularization, reducing model complexity, or gathering more data."
        )

    doc.add_paragraph(conclusion)

    # --- Save ---
    today = date.today().strftime("%Y%m%d")
    safe_name = profile.filename.replace(".", "_")
    filename = f"ml_report_{safe_name}_{today}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    return os.path.abspath(filepath)
